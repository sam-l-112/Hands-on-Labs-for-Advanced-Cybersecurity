#!/usr/bin/env python3
"""Generate W13_Code_Smell_學號_姓名.docx from the lab report."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft JhengHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

def add_heading_t(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft JhengHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    return h

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    run.font.name = 'Microsoft JhengHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Microsoft JhengHei'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Microsoft JhengHei'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    return table

def add_screenshot_placeholder(label):
    p = doc.add_paragraph()
    run = p.add_run(f'\n【{label}】請在此插入截圖\n')
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(10)
    run.font.name = 'Microsoft JhengHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

# ===== Title =====
title = doc.add_heading('安全程式設計 × AI Agent 協作實驗報告', level=0)
for run in title.runs:
    run.font.name = 'Microsoft JhengHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

add_para('')

add_para('實驗日期：2026/05/18–05/24')
add_para('姓名：', bold=True)
add_para('（請填寫）', italic=True)
add_para('學號：', bold=True)
add_para('（請填寫）', italic=True)

doc.add_page_break()

# ===== 一、執行摘要 =====
add_heading_t('一、執行摘要', level=1)
add_para(
    '本次實驗透過三個 Lab 驗證 AI agent（opencode）在安全程式設計中的能力邊界。'
    '最重要的發現是：功能測試（TDD）全數通過不等於程式安全——'
    'Lab 1 中使用 gets() 與 strcmp() 的程式碼能通過所有 assert，'
    '但存在 4 項 CERT C 違規。此外，opencode 對訓練資料截止日期後的未知 CVE'
    '（Dirty Frag CVE-2026-43284）仍能透過歷史相似模式正確回答，'
    '但這不代表 AI 的判斷可以完全取代人類驗證。'
    '整體而言，AI agent 是有記憶截止日期的協作者，指引品質直接決定輸出安全等級。'
)

doc.add_page_break()

# ===== 二、Lab 1 =====
add_heading_t('二、Lab 1 — TDD × opencode × 密碼驗證', level=1)

add_heading_t('2.1 TDD Red：測試截圖', level=2)
add_screenshot_placeholder('截圖 1-A：opencode 生成的 assert 測試程式碼')

add_heading_t('2.2 TDD Green：實作截圖', level=2)
add_screenshot_placeholder('截圖 1-B：opencode 依 gets() / strcmp() 實作的 verify_password()')

# 2.3 Code Smell 清單
add_heading_t('2.3 Code Smell 清單', level=2)

headers_smell = [
    '#', '程式碼位置', 'Smell 描述', 'CERT 規則',
    'CWE 編號', '若不修正，最壞後果', '功能測試有抓到嗎？'
]
rows_smell = [
    ['1', 'verify_password() — gets(buf)',
     '使用 gets() 無長度限制，輸入超過 19 字元即 buffer overflow',
     'STR31-C / STR07-C', 'CWE-120 / CWE-676',
     '攻擊者可覆蓋 return address 執行任意程式碼', '❌ 否'],
    ['2', 'verify_password() — strcmp(buf, PASSWORD)',
     '明文比對密碼，存在 timing attack 風險',
     'MSC41-C', 'CWE-798',
     '攻擊者可透過回應時間差異逐步猜出密碼', '❌ 否'],
    ['3', 'verify_password() — #define MAX_BUF 20',
     'magic number 寫死 buffer 大小，密碼長度接近 buffer 上限',
     'STR31-C', 'CWE-120',
     '密碼若超過 19 字元即溢位', '❌ 否'],
    ['4', 'verify_password() — 未檢查 gets() 回傳值',
     '未確認 gets() 回傳值是否為 NULL（EOF 時回傳 NULL）',
     'ERR33-C', 'CWE-391',
     'EOF 時進入未定義行為', '❌ 否'],
]
add_table(headers_smell, rows_smell)

add_para('')
add_para('觀察：最後一欄「功能測試有抓到嗎？」全部填完後，你發現了什麼規律？', bold=True)
add_para(
    '功能測試（assert）的設計目標是驗證「行為是否正確」，而非「實作是否安全」。'
    '所有 smell 都是實作品質問題而非功能錯誤，因此功能測試完全無法偵測 security smell。'
    '這說明了安全測試需要獨立設計，不能依附在功能測試之下。'
)

# 2.4 opencode 自審結果
add_heading_t('2.4 opencode 自審結果', level=2)
add_screenshot_placeholder('截圖 1-C：opencode 自審的輸出')

add_para('三方比較', bold=True)
headers_3way = ['Smell', '功能測試（assert）抓到？', 'opencode 自審抓到？', '你用 CERT C 掃到？']
rows_3way = [
    ['gets() buffer overflow', '❌', '✅ 有指出 STR31-C', '✅ STR31-C / STR07-C'],
    ['strcmp() timing attack', '❌', '⚠️ 有提但未深入細節', '✅ MSC41-C'],
    ['寫死的常數', '❌', '✅ 有指出 MSC41-C', '✅ CWE-798'],
    ['錯誤訊息洩漏', '❌', '⚠️ 僅提醒未來注意', '✅ ERR07-C / CWE-209'],
]
add_table(headers_3way, rows_3way)

add_para('')
p = add_para('核心結論：', bold=True)
add_para(
    '不能。功能測試是針對「行為」設計的：只要回傳值正確，assert 就是綠燈。'
    '而 security smell 是「實作品質」的問題——程式可能功能完全正確，'
    '但寫法上存在遲早會被利用的弱點。'
    '安全審查需要獨立的檢查維度（如 CERT C 規則、威脅模型、CWE 分類），'
    '與功能測試是互補關係而非替代關係。'
)

doc.add_page_break()

# ===== 三、Lab 2 =====
add_heading_t('三、Lab 2 — opencode 分析未知 CVE 片段', level=1)

add_heading_t('3.1 Q1：功能描述', level=2)
add_screenshot_placeholder('截圖 2-A：opencode 對 Q1 的回答')

add_para('驗證結果', bold=True)
headers_q1 = ['評估項目', '是 / 否', '說明']
rows_q1 = [
    ['正確描述了 skb fragment 處理', '是',
     '正確指出 skb_shinfo(skb)->frags 走訪與 kmap_atomic + memset 行為'],
    ['提到了 page cache 或 shared memory 的概念', '是',
     '有提及「頁面映射到核心虛擬空間」與「同一塊記憶體可能被其他執行緒存取」'],
    ['解釋清楚讓非 kernel 工程師也能懂', '是',
     '使用「清理網路封包碎片」的比喻，易懂'],
]
add_table(headers_q1, rows_q1)

add_heading_t('3.2 Q2：資安問題', level=2)
add_screenshot_placeholder('截圖 2-B：opencode 對 Q2 的回答')

add_para('驗證結果', bold=True)
headers_q2 = ['評估項目', '是 / 否', '說明']
rows_q2 = [
    ['指出了 race condition / TOCTOU 性質的問題', '是',
     '明確指出「無加鎖 → 競爭條件 → TOCTOU」的因果鏈'],
    ['提到 COW（Copy-On-Write）缺失', '是',
     '指出可被用於 COW bypass，繞過記憶體隔離'],
    ['對應到 CERT CON30-C 或 FIO45-C', '是',
     '同時對應 CON30-C（共享資源未加鎖）與 FIO45-C（TOCTOU）'],
    ['說明了漏洞可能的攻擊情境', '是',
     '說明攻擊者可讓核心在清空頁面前寫入惡意資料以繞過 COW'],
]
add_table(headers_q2, rows_q2)

add_heading_t('3.3 Q3：CVE 編號', level=2)
add_screenshot_placeholder('截圖 2-C：opencode 對 Q3 的回答')

add_para('NVD 驗證', bold=True)
add_para(
    '（請至 https://nvd.nist.gov 搜尋 CVE-2026-43284，將 NVD 頁面截圖貼在此處）',
    italic=True
)

headers_cve = ['評估項目', '結果']
rows_cve = [
    ['opencode 回答的 CVE 編號', 'CVE-2026-43284'],
    ['正確 CVE 編號', 'CVE-2026-43284'],
    ['opencode 答對了嗎', '✅ 是'],
    ['若答錯，opencode 說了什麼', '不適用（回答正確）'],
]
add_table(headers_cve, rows_cve)

add_heading_t('3.4 反思', level=2)

add_para('思考題 1：opencode 對 Q3 的表現說明了什麼？是 AI 不夠聰明、還是有別的原因？', bold=True)
add_para(
    'opencode 能正確回答 CVE-2026-43284，並非因為它「知道」這個 2026/05/07 才公開的漏洞，'
    '而是因為 Dirty Frag 的底層模式——page cache race condition 導致 COW bypass——'
    '與歷史上的 Dirty COW（CVE-2016-5195）幾乎相同。'
    'opencode 的訓練資料中有大量 COW bypass 的案例，因此它能透過模式匹配推斷出正確的 CVE 編號。'
    '這說明 AI 的能力來自於訓練資料的覆蓋範圍與相似模式的遷移學習，而不是真正的「漏洞知識」。'
    '關鍵教訓：AI 的答案始終需要你對照 NVD、CERT 等權威來源來驗證，不應盲目信任。'
)

add_para('思考題 2：如果你在做期末滲透測試時，目標系統有一個 2026 年才公開的漏洞，你應該怎麼做？', bold=True)
add_para(
    '首先，不該僅依賴 AI agent 的知識——它的訓練資料有截止日期，對全新的漏洞可能存在盲區。'
    '其次，應主動查閱權威來源：NVD（nvd.nist.gov）、OSS-SEC mailing list、廠商安全公告、'
    'GitHub Advisory Database、Twitter 安全研究員 timeline。'
    '第三，一旦從上述管道確認漏洞存在，應將漏洞細節作為指引提供給 opencode，'
    '讓它協助撰寫測試程式碼或 exploit 雛形，而非讓它自行猜測 CVE 編號。'
    '最後，使用自動化工具（如 semgrep、metasploit、Nessus）交叉驗證，確保覆蓋率。'
)

doc.add_page_break()

# ===== 四、Lab 3 =====
add_heading_t('四、Lab 3 — TDD × 好指引下的重寫', level=1)

add_heading_t('4.1 TDD Red：含安全測試的測試集截圖', level=2)
add_screenshot_placeholder('截圖 3-A：含功能測試 + 安全測試 comment 的測試集')

add_heading_t('4.2 我的「好指引」', level=2)
add_screenshot_placeholder('截圖 3-B：完整好指引')

add_para('好指引包含的元素確認：', bold=True)
headers_guide = ['元素', '是否包含', '簡述你寫了什麼']
rows_guide = [
    ['功能 spec', '✅ 是',
     '從 stdin 讀取密碼，與預設密碼比對，正確回傳 1 錯誤回傳 0'],
    ['Threat model', '✅ 是',
     'Timing attack（回應時間差異猜密碼）、超長輸入 crash、null bytes 處理'],
    ['不可接受清單（禁止 gets / strcmp）', '✅ 是',
     '禁止 gets()（STR31-C）、禁止 strcmp()（timing leak）、禁止明文密碼（MSC41-C）'],
    ['安全測試案例（超長輸入、timing）', '✅ 是',
     '測試 4（10000 字元不 crash）、測試 5（100 次呼叫 timing 一致性）'],
]
add_table(headers_guide, rows_guide)

add_heading_t('4.3 Lab 1 vs Lab 3 對比', level=2)
headers_cmp = ['Smell 類型', 'Lab 1 有', 'Lab 3 有', 'Lab 3 的哪個安全測試防住了它']
rows_cmp = [
    ['危險輸入函式（gets）', '✅ 有', '❌ 無',
     '測試 4（超長輸入不 crash）強迫使用 fgets 而非 gets'],
    ['明文比對（strcmp）/ timing attack', '✅ 有', '❌ 無',
     '測試 5（timing 一致性）強迫使用 constant_time_memcmp'],
    ['寫死的常數', '✅ 有', '⚠️ hash 仍寫死',
     '指引禁止明文密碼，但 SHA-256 hash 仍寫在程式碼中'],
    ['缺少長度限制', '✅ 有', '❌ 無',
     '測試 4（10000 字元）確保 buffer 256 + 截斷機制'],
    ['錯誤訊息洩漏', '✅ 有', '❌ 無',
     '指引第 6 條明確禁止洩漏內部路徑'],
]
add_table(headers_cmp, rows_cmp)

add_para('')
add_para('觀察：加入安全測試後，opencode 有沒有主動避開對應的 smell？', bold=True)
add_para(
    '加入安全測試後，opencode 在 5 項 smell 中有 4 項成功避開。'
    '關鍵在於安全測試「約束了實作的邊界條件」——例如測試 4（10000 字元不 crash）'
    '直接排除了 gets() 的可能性，因為 gets() 無法處理超長輸入。'
    '而測試 5（timing 一致性）強迫使用常數時間比較，避免了 strcmp()。'
    '這證明了安全測試案例不僅是驗證工具，更是設計約束，'
    '能有效引導 AI agent 避開已知的危險模式。'
)

doc.add_page_break()

# ===== 五、整體反思 =====
add_heading_t('五、整體反思', level=1)

add_para('最讓你印象深刻的發現是什麼？', bold=True)
add_para(
    '最令我印象深刻的是 Lab 1 中「測試全過但程式不安全」的矛盾現象。'
    '三個 assert 全部綠燈，但程式碼卻存在 buffer overflow、timing attack、'
    '硬編碼密碼等 4 項嚴重 security smell。'
    '這徹底打破了「TDD 通過 = 程式碼品質好」的直覺。'
    '進一步對比 Lab 3 可以發現，同樣是 TDD，只要在測試集中加入安全屬性'
    '（超長輸入不 crash、timing 一致性），opencode 的輸出品質就有顯著提升。'
    '這說明了指引（prompt）才是控制 AI 輸出品質的核心槓桿——'
    '你給的條件越精確，AI 的表現就越好。'
)

add_para('你認為「指引品質就是你的能力」這個說法成立嗎？為什麼？', bold=True)
add_para(
    '成立。Lab 1 與 Lab 3 的強烈對比完全支持這個觀點。'
    '同樣的 AI agent（opencode）、同樣的任務（寫密碼驗證函式），'
    '僅僅因為指引內容不同，就產生了「含有 4 個 CVE 等級漏洞的程式碼」'
    'vs 「符合 CERT C 規則的安全程式碼」的天壤之別。'
    '好指引的關鍵元素是：明確的功能 spec + 威脅模型 + 不可接受清單 + 安全測試案例。'
    '這四者缺一不可。換句話說，AI agent 只是一個執行引擎，'
    '真正決定輸出品質的是人類賦予它的「思考框架」。'
    '能寫出好指引的人，才真正掌控了 AI 輔助開發的安全與品質。'
)

doc.add_page_break()

# ===== 六、時事連結 =====
add_heading_t('六、時事連結', level=1)

add_para('選項 B：Dirty Frag（CVE-2026-43284）的 COW bypass 模式，跟 Week 12 的 Pack2TheRoot TOCTOU 有什麼共同點？從 CERT C 規則的角度，用 CON30-C 和 FIO45-C 各解釋一次。', bold=True)
add_para('')
add_para('這兩個漏洞的共同點是「在 check 與 use 之間，狀態被非法改變」。', bold=True)
add_para('')
add_para('CON30-C（共享資源未加鎖）：', bold=True)
add_para(
    'Dirty Frag 的 fragment_process() 在走訪 skb fragment 時，'
    '對每個 page 進行 kmap_atomic() + memset()，但整個過程未使用任何鎖機制。'
    '另一個執行緒可以同時修改同一個 page 的內容，'
    '導致核心在 memset 清空頁面時，攻擊者已經寫入了惡意資料，達成 COW bypass。'
    'Pack2TheRoot 中，Checker 模組檢查檔案路徑時未對共享的目錄結構加鎖，'
    '導致在檢查通過後、實際操作前，符號連結已被更換。'
    '兩者都是 CON30-C 違規——對共享資源的存取缺乏同步機制。'
)
add_para('')
add_para('FIO45-C（避免 TOCTOU）：', bold=True)
add_para(
    'Dirty Frag 中，核心讀取 skb_shinfo(skb)->nr_frags 與 skb_frag_page(frag) 之後，'
    '到實際 kmap_atomic() + memset() 之間，fragment 的底層 page 可能已被釋放或替換——'
    '典型的 TOCTOU。'
    'Pack2TheRoot 中，access(path, R_OK) 檢查檔案可讀性之後、open(path) 之前，'
    '攻擊者已將路徑指向敏感檔案（符號連結替換）。'
    '兩者都是 FIO45-C 違規——檢查（check）與使用（use）之間的時間窗口未被消除。'
)
add_para('')
add_para('根本差異：', bold=True)
add_para(
    'Dirty Frag 的 TOCTOU 發生在記憶體層級（page cache 的並行寫入），'
    'Pack2TheRoot 的 TOCTOU 發生在檔案系統層級（符號連結替換）。'
    '但解決方案一致：要麼在 check-use 之間加鎖（CON30-C），'
    '要麼使用原子操作消除時間窗口（FIO45-C）。'
)

# ===== Save =====
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'W13_Code_Smell_學號_姓名.docx'
)
doc.save(output_path)
print(f'Document saved to: {output_path}')
